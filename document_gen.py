from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class DocumentGenerator:
    def __init__(self):
        self.output_dir = "generated_docs"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def create_document(self, content: str, title: str = "Generated Document") -> str:
        """Create a Word document from content"""
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add separator
        doc.add_paragraph("_" * 60)
        
        # Process content into sections
        sections = content.split("\n\n")
        
        for section in sections:
            if section.strip():
                # Check if it looks like a heading
                if section.startswith("#") or len(section) < 80:
                    # Clean heading
                    heading_text = section.replace("#", "").strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=1)
                else:
                    # Regular paragraph
                    p = doc.add_paragraph(section)
                    p.style.font.size = Pt(11)
        
        # Save document
        filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath